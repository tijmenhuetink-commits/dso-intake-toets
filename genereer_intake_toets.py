"""
DSO Intake Toets Generator
===========================
Versie : 3.3
Datum  : 2026-09-03
Wijzigingen:
  v1.0 — eerste versie, Word-document gegenereerd vanuit DSO JSON-data
  v2.0 — niet-gedigitaliseerde plannen netjes afgehandeld
  v3.0 — volledig nieuw: gebaseerd op officieel ODU-sjabloon
  v3.1 — voorbereidingsbesluit-filter: alleen echte VB-plannen tonen op basis van plan-ID
          oranje tekst voor automatisch ingevulde DSO-velden
  v3.2 — RD-coördinaten ingevuld in Locatie-tabel
  v3.3 — maatvoeringen gededupliceerd; hyperlink/omgevingsplan-cel fix
          gebiedsaanduidingen ingevuld als kommalijst
          sjabloon bijgewerkt naar Intaketoets.docx (Intaketoets.docx)
          exact dezelfde opmaak als het sjabloon, geen kleurcoderingen
          DSO-velden automatisch ingevuld op de juiste plekken
          Toets Omgevingsplan tekst automatisch samengesteld

Genereert een ingevuld Word-document (Intake_toets) op basis van de
JSON-output van dso_bestemmingsplan.py, exact in het formaat van het
ODU-sjabloon.

Gebruik:
  python genereer_intake_toets.py                        # vraagt adres interactief
  python genereer_intake_toets.py "Kerkstraat 1, Utrecht"
  python genereer_intake_toets.py 131653,447223          # RD-coördinaten

Output:
  Intake_toets_<adres>_<datum>.docx  op de Desktop

Benodigdheden:
  pip install requests python-docx
"""

VERSION = "3.3"

import sys
import os
import json
import copy
import importlib.util
from datetime import date
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor


# ─────────────────────────────────────────────
# SJABLOON PAD
# ─────────────────────────────────────────────
eigen_map = os.path.dirname(os.path.abspath(sys.argv[0]))
SJABLOON_PAD = os.path.join(eigen_map, "Intaketoets.docx")


# ─────────────────────────────────────────────
# DSO-SCRIPT LADEN
# ─────────────────────────────────────────────
def _laad_dso():
    for pad in [
        os.path.join(eigen_map, "dso_bestemmingsplan.py"),
        os.path.join(os.getcwd(), "dso_bestemmingsplan.py"),
    ]:
        if os.path.isfile(pad):
            spec = importlib.util.spec_from_file_location("dso_bestemmingsplan", pad)
            mod  = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            return mod
    return None

_dso = _laad_dso()
if _dso:
    haal_data_voor_adres       = _dso.haal_data_voor_adres
    haal_data_voor_coordinaten = _dso.haal_data_voor_coordinaten
    DSO_BESCHIKBAAR = True
else:
    DSO_BESCHIKBAAR = False
    print("⚠  dso_bestemmingsplan.py niet gevonden")


# ─────────────────────────────────────────────
# HULPFUNCTIES
# ─────────────────────────────────────────────

def is_gemeentelijk_plan_id(plan_id):
    """Check of plan-ID een gemeentelijk plan is (niet rijks/provinciaal)."""
    if not plan_id.startswith("NL.IMRO."):
        return False
    rest = plan_id[len("NL.IMRO."):]
    gemeentecode = rest[:4]
    return gemeentecode.isdigit() and gemeentecode != "0000"


def is_echt_vbb(plan):
    """Check of een plan echt een voorbereidingsbesluit is op basis van plan-ID."""
    plan_id = plan.get("id", "")
    if not plan_id.startswith("NL.IMRO."):
        return False
    rest = plan_id[len("NL.IMRO."):][4:].lstrip(".").upper()
    return rest.startswith("VB") or rest.startswith("VRO") or "VROBB" in rest


def para_tekst(para):
    """Geeft de volledige tekst van een paragraaf terug."""
    return ''.join(r.text or '' for r in para._p.iter(qn('w:t')))


def zoek_cel_naast(doc, zoektekst):
    """
    Zoekt in alle tabellen naar een cel met zoektekst en geeft
    de cel rechts ernaast terug (de waardecel).
    """
    for tbl in doc.tables:
        for rij in tbl.rows:
            for i, cel in enumerate(rij.cells):
                if zoektekst.lower() in cel.text.lower():
                    if i + 1 < len(rij.cells):
                        return rij.cells[i + 1]
    return None


# Donker oranje kleur voor automatisch ingevulde DSO-velden
ORANJE = RGBColor(0xC0, 0x50, 0x00)  # donker oranje


def vul_cel(cel, tekst, hyperlink_url=None, oranje=False):
    """
    Vult een tabelcel met tekst, behoudt bestaande opmaak.
    Als hyperlink_url opgegeven, maakt een klikbare link.
    Als oranje=True, wordt de tekst donker oranje (automatisch ingevuld via DSO).
    """
    if cel is None:
        return

    # Verwijder bestaande inhoud maar bewaar de paragraaf
    for para in cel.paragraphs:
        for run in para.runs:
            run.text = ''
        # Verwijder eventuele hyperlinks
        for child in list(para._p):
            if child.tag == qn('w:hyperlink'):
                para._p.remove(child)

    if not cel.paragraphs:
        cel.add_paragraph()

    para = cel.paragraphs[0]

    if hyperlink_url:
        _voeg_hyperlink_toe(para, tekst, hyperlink_url)
    else:
        run = para.add_run(tekst)
        if oranje:
            run.font.color.rgb = ORANJE
        # Kopieer opmaak van eventueel eerste bestaande run
        if len(para.runs) > 1:
            eerste = para.runs[0]
            run.bold = eerste.bold
            run.font.size = eerste.font.size
            run.font.name = eerste.font.name


def _voeg_hyperlink_toe(para, tekst, url):
    """Voegt een klikbare hyperlink toe aan een paragraaf."""
    # Registreer de relatie
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    # Maak het hyperlink XML-element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    # Hyperlink stijl (blauw, onderstreept)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rpr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rpr.append(u)
    run_elem.append(rpr)

    t = OxmlElement('w:t')
    t.text = tekst
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


def vul_para_met_tekst(para, tekst):
    """Vervangt de tekst van een paragraaf, behoudt opmaak van eerste run."""
    # Bewaar opmaak
    rpr_xml = None
    if para.runs:
        rpr_xml = para.runs[0]._r.find(qn('w:rPr'))

    # Leeg de paragraaf
    for run in para.runs:
        run.text = ''
    # Verwijder hyperlinks
    for child in list(para._p):
        if child.tag == qn('w:hyperlink'):
            para._p.remove(child)

    # Voeg nieuwe tekst toe
    run = para.add_run(tekst)
    if rpr_xml is not None:
        run._r.insert(0, copy.deepcopy(rpr_xml))


# ─────────────────────────────────────────────
# HOOFDFUNCTIE
# ─────────────────────────────────────────────

def genereer_intake_toets(data: dict, uitvoer_pad: str = None) -> str:
    """
    Genereert een ingevuld Intake Toets document op basis van DSO-data.
    Gebruikt het officiële ODU-sjabloon als basis.

    Args:
        data: dict met DSO-data (output van dso_bestemmingsplan.py)
        uitvoer_pad: pad waar het document opgeslagen wordt (optioneel)

    Returns:
        pad naar het gegenereerde document
    """
    # Controleer of sjabloon bestaat
    if not os.path.isfile(SJABLOON_PAD):
        raise FileNotFoundError(
            f"Sjabloon niet gevonden: {SJABLOON_PAD}\n"
            f"Zorg dat 'Intaketoets.docx' in dezelfde map staat als dit script."
        )

    # Open het sjabloon
    doc = Document(SJABLOON_PAD)

    # ── Data ophalen ──────────────────────────────────────────────────────────
    adres_gevonden  = data.get("adres_gevonden") or data.get("adres", "—")
    bp_naam         = data.get("bestemmingsplan_naam", "—")
    bp_datum        = data.get("bestemmingsplan_datum", "—")
    hyperlink       = data.get("hyperlink", "—")
    bestemming      = data.get("bestemming_perceel", "—")
    bestemmingstype = data.get("bestemmingstype", "—")
    kadaster        = data.get("kadastrale_aanduiding", "—")
    niet_gedig      = data.get("niet_gedigitaliseerd", False)

    functie_str = ", ".join(data.get("functieaanduidingen", [])) or "geen"
    dubbel_str  = ", ".join(d["naam"] for d in data.get("dubbelbestemmingen", [])) or "geen"
    bouw_str    = ", ".join(b["naam"] for b in data.get("bouwaanduidingen", [])) or "geen"

    maatvoeringen = data.get("maatvoeringen", [])
    # Dedupliceer maatvoeringen op naam (bewaar eerste unieke waarde)
    gezien_maat = {}
    for m in maatvoeringen:
        naam = m.get("naam", "").lower()
        if naam not in gezien_maat:
            gezien_maat[naam] = m
    maatvoeringen_uniek = list(gezien_maat.values())

    def maatv(zoektermen):
        for z in zoektermen:
            for m in maatvoeringen_uniek:
                if z in m.get("naam", "").lower():
                    return f"{m['waarde']} {m.get('eenheid', '')}".strip()
        return "—"

    bouwhoogte = maatv(["bouwhoogte"])
    goothoogte = maatv(["goothoogte"])
    opp        = maatv(["oppervlakte", "bebouwd", "bouwperceel"])

    # Planenoverzicht
    planenoverzicht = data.get("planenoverzicht", {})
    vbb_str = ", ".join(
        p.get("naam", "—") for p in planenoverzicht.get("voorbereidingsbesluit", [])
        if is_echt_vbb(p) and is_gemeentelijk_plan_id(p.get("id", ""))
    ) or "—"
    parkeer_str = "—"
    for p in planenoverzicht.get("bestemmingsplan", []):
        if "parkeer" in p.get("naam", "").lower() and p.get("paraplu"):
            parkeer_str = p.get("naam", "—")
            break

    # Coördinaten
    coord_x = data.get("x")
    coord_y = data.get("y")
    coord_str = f"X={coord_x:.2f}, Y={coord_y:.2f}" if coord_x and coord_y else "—"

    # Gebiedsaanduidingen
    gebiedsaand_str = ", ".join(
        g.get("naam", "—") for g in data.get("alle_gebiedsaanduidingen", [])
    ) or "—"

    # Adres splitsen
    adres_delen = adres_gevonden.split(", ") if ", " in adres_gevonden else [adres_gevonden]
    straat_hnr  = adres_delen[0]

    # ── Velden invullen ───────────────────────────────────────────────────────

    # Locatie: Straatnaam + huisnummer
    cel = zoek_cel_naast(doc, "Straatnaam + huisnummer")
    vul_cel(cel, straat_hnr, oranje=True)

    # Locatie: RD-coördinaten
    cel = zoek_cel_naast(doc, "RD-coördinaten (ingevoerd)")
    vul_cel(cel, coord_str, oranje=True)

    # Hyperlink regels op de kaart — zoek specifiek in tabel 6 (Omgevingsplan informatie)
    for tbl in doc.tables:
        for rij in tbl.rows:
            if "Hyperlink regels op de kaart" in rij.cells[0].text:
                if len(rij.cells) > 1:
                    vul_cel(rij.cells[1], "Link",
                        hyperlink_url="https://omgevingswet.overheid.nl/regels-op-de-kaart/zoeken/locatie")
                break

    # Omgevingsplan — zoek alleen de rij die direct na "Hyperlink regels op de kaart" staat
    gevonden_hyperlink = False
    for tbl in doc.tables:
        for rij in tbl.rows:
            if gevonden_hyperlink and rij.cells[0].text.strip() == "Omgevingsplan":
                vul_cel(rij.cells[1], "zie Regels op de kaart")
                gevonden_hyperlink = False
                break
            if "Hyperlink regels op de kaart" in rij.cells[0].text:
                gevonden_hyperlink = True

    # Bestemmingsplan
    cel = zoek_cel_naast(doc, "Bestemmingsplan")
    vul_cel(cel, f"{bp_naam}" + (f" ({bp_datum})" if bp_datum and bp_datum != "—" else ""), oranje=True)

    # Bestemming perceel
    cel = zoek_cel_naast(doc, "Bestemming perceel")
    vul_cel(cel, bestemming if not niet_gedig else "zie hyperlink plan", oranje=True)

    # Dubbelbestemming
    cel = zoek_cel_naast(doc, "Dubbelbestemming")
    vul_cel(cel, dubbel_str if not niet_gedig else "zie hyperlink plan", oranje=True)

    # (Functie)aanduiding
    cel = zoek_cel_naast(doc, "(Functie)aanduiding")
    vul_cel(cel, functie_str if not niet_gedig else "zie hyperlink plan", oranje=True)

    # Gebiedsaanduiding
    cel = zoek_cel_naast(doc, "Gebiedsaanduiding")
    vul_cel(cel, gebiedsaand_str, oranje=True)

    # Voorbereidingsbesluit
    cel = zoek_cel_naast(doc, "Voorbereidingsbesluit")
    vul_cel(cel, vbb_str)

    # Bestemmingsplan parkeren
    cel = zoek_cel_naast(doc, "Bestemmingsplan parkeren")
    vul_cel(cel, parkeer_str)

    # Bebouwde oppervlakte
    cel = zoek_cel_naast(doc, "Bebouwde oppervlakte bouwperceel")
    vul_cel(cel, opp if opp != "—" else "niet opgenomen in plan", oranje=True)

    # Maximale bouwhoogte
    cel = zoek_cel_naast(doc, "Maximale bouwhoogte")
    vul_cel(cel, bouwhoogte if bouwhoogte != "—" else "niet opgenomen in plan", oranje=True)

    # Maximale goothoogte
    cel = zoek_cel_naast(doc, "Maximale goothoogte")
    vul_cel(cel, goothoogte if goothoogte != "—" else "niet opgenomen in plan", oranje=True)

    # ── Toets Omgevingsplan tekst ─────────────────────────────────────────────
    # Zoek de paragraaf met <GLOBALE LOCATIE>
    for para in doc.paragraphs:
        tekst = para_tekst(para)
        if "<GLOBALE LOCATIE>" in tekst or "GLOBALE LOCATIE" in tekst:
            # Stel de samengestelde tekst samen
            dubbel_deel = f", met de dubbelbestemming '{dubbel_str}'" if dubbel_str != "geen" else ""
            bouw_deel   = f" en bouwaanduiding '{bouw_str}'" if bouw_str != "geen" else ""
            functie_deel = f" en functieaanduiding '{functie_str}'" if functie_str != "geen" else ""
            parkeer_deel = f" Daarnaast is het bestemmingsplan parkeren '{parkeer_str}' van toepassing." if parkeer_str != "—" else ""

            nieuwe_tekst = (
                f"Het perceel {straat_hnr} ligt binnen het plangebied van het vigerende "
                f"bestemmingsplan \"{bp_naam}\" (vastgesteld {bp_datum}). "
                f"Het perceel heeft op basis daarvan de {bestemmingstype} '{bestemming}'"
                f"{dubbel_deel}{bouw_deel}{functie_deel}."
                f"{parkeer_deel} "
                f"De aanvraag omgevingsvergunning is gesitueerd op deze bestemming."
            )
            vul_para_met_tekst(para, nieuwe_tekst)
            break

    # ── Opslaan ───────────────────────────────────────────────────────────────
    if uitvoer_pad is None:
        desktop   = os.path.join(os.path.expanduser("~"), "Desktop")
        adres_kort = straat_hnr.replace(" ", "_").replace("/", "-")[:25]
        bestandsnaam = f"Intake_toets_{adres_kort}_{date.today().strftime('%Y%m%d')}.docx"
        uitvoer_pad = os.path.join(desktop, bestandsnaam)

    doc.save(uitvoer_pad)
    print(f"\n✓ Word-document opgeslagen: {uitvoer_pad}")
    return uitvoer_pad


# ─────────────────────────────────────────────
# UITVOEREN
# ─────────────────────────────────────────────

if __name__ == "__main__":

    def vraag_invoer():
        print("=" * 55)
        print("  DSO Intake Toets Generator  v" + VERSION)
        print("=" * 55)
        print("  Kies invoermethode:")
        print("  1. Adres (bijv. Kerkstraat 1, IJsselstein)")
        print("  2. RD-coördinaten (bijv. 131653, 447223)")
        print("  3. Testadres (Graaf Walramhof 4, Nieuwegein)")
        print("=" * 55)
        keuze = input("  Keuze [1/2/3]: ").strip()
        if keuze == "2":
            coords = input("  X, Y (RD): ").strip()
            try:
                x_str, y_str = coords.replace(" ", "").split(",")
                return None, float(x_str), float(y_str)
            except Exception:
                return "Graaf Walramhof 4, Nieuwegein", None, None
        elif keuze == "3":
            return "Graaf Walramhof 4, Nieuwegein", None, None
        else:
            adres = input("  Adres: ").strip()
            return (adres or "Graaf Walramhof 4, Nieuwegein"), None, None

    if not DSO_BESCHIKBAAR:
        print("✗ dso_bestemmingsplan.py niet gevonden. Zet beide scripts in dezelfde map.")
        sys.exit(1)

    if len(sys.argv) > 1:
        arg = " ".join(sys.argv[1:])
        if "," in arg and arg.replace(",","").replace(".","").replace(" ","").replace("-","").isdigit():
            parts = arg.replace(" ","").split(",")
            invoer_adres, invoer_x, invoer_y = None, float(parts[0]), float(parts[1])
        else:
            invoer_adres, invoer_x, invoer_y = arg, None, None
    else:
        invoer_adres, invoer_x, invoer_y = vraag_invoer()

    try:
        if invoer_x is not None:
            data = haal_data_voor_coordinaten(invoer_x, invoer_y)
        else:
            data = haal_data_voor_adres(invoer_adres)
        genereer_intake_toets(data)
    except FileNotFoundError as e:
        print(f"\n✗ {e}")
    except Exception as e:
        print(f"\n✗ Fout: {e}")
        raise
