"""
DSO Intake Toets — Streamlit Web App
=====================================
Draait op Streamlit Community Cloud.
Vereist: streamlit, requests, python-docx
"""

import streamlit as st
import requests
import json
import re
from datetime import date
from io import BytesIO

# ── Pagina-instellingen ───────────────────────────────────────────────────────
st.set_page_config(
    page_title="DSO Intake Toets",
    page_icon="🏛️",
    layout="centered",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'IBM Plex Sans', sans-serif;
}

/* Achtergrond */
.stApp {
    background: #0d1117;
    color: #e6edf3;
}

/* Header blok */
.dso-header {
    background: linear-gradient(135deg, #1f4e79 0%, #0d2137 100%);
    border: 1px solid #1f6feb;
    border-radius: 8px;
    padding: 2rem 2.5rem;
    margin-bottom: 2rem;
}
.dso-header h1 {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 1.6rem;
    font-weight: 600;
    color: #58a6ff;
    margin: 0 0 0.3rem 0;
    letter-spacing: -0.5px;
}
.dso-header p {
    color: #8b949e;
    font-size: 0.9rem;
    margin: 0;
}

/* Stap-labels */
.stap-label {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.7rem;
    color: #58a6ff;
    text-transform: uppercase;
    letter-spacing: 2px;
    margin-bottom: 0.3rem;
}

/* Resultaat kaarten */
.result-card {
    background: #161b22;
    border: 1px solid #30363d;
    border-radius: 6px;
    padding: 1.2rem 1.5rem;
    margin: 0.5rem 0;
}
.result-card .label {
    font-size: 0.75rem;
    color: #8b949e;
    font-family: 'IBM Plex Mono', monospace;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin-bottom: 0.2rem;
}
.result-card .waarde {
    font-size: 1rem;
    color: #e6edf3;
    font-weight: 600;
}
.result-card .waarde.auto {
    color: #58a6ff;
}
.result-card .waarde.leeg {
    color: #484f58;
    font-style: italic;
    font-weight: 400;
}

/* Terminal output */
.terminal {
    background: #0d1117;
    border: 1px solid #21262d;
    border-radius: 6px;
    padding: 1rem 1.2rem;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.8rem;
    color: #3fb950;
    white-space: pre-wrap;
    max-height: 300px;
    overflow-y: auto;
}

/* Legenda badges */
.badge-auto { background: #1f3a5f; color: #58a6ff; padding: 2px 10px; border-radius: 20px; font-size: 0.75rem; margin-right: 6px; }
.badge-hand { background: #3d2e00; color: #d29922; padding: 2px 10px; border-radius: 20px; font-size: 0.75rem; margin-right: 6px; }
.badge-pb   { background: #3d1a1a; color: #f85149; padding: 2px 10px; border-radius: 20px; font-size: 0.75rem; }

/* Download knop */
.stDownloadButton > button {
    background: #238636 !important;
    color: white !important;
    border: 1px solid #2ea043 !important;
    border-radius: 6px !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
    font-weight: 600 !important;
    padding: 0.6rem 1.5rem !important;
    width: 100% !important;
    font-size: 1rem !important;
}
.stDownloadButton > button:hover {
    background: #2ea043 !important;
}

/* Input veld */
.stTextInput > div > div > input {
    background: #161b22 !important;
    border: 1px solid #30363d !important;
    color: #e6edf3 !important;
    border-radius: 6px !important;
    font-family: 'IBM Plex Mono', monospace !important;
    font-size: 1rem !important;
}
.stTextInput > div > div > input:focus {
    border-color: #58a6ff !important;
    box-shadow: 0 0 0 3px rgba(88,166,255,0.1) !important;
}

/* Primary button */
.stButton > button[kind="primary"] {
    background: #1f4e79 !important;
    color: white !important;
    border: 1px solid #1f6feb !important;
    border-radius: 6px !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
    font-weight: 600 !important;
    width: 100% !important;
    font-size: 1rem !important;
    padding: 0.6rem 1.5rem !important;
}
.stButton > button[kind="primary"]:hover {
    background: #2d6da8 !important;
    border-color: #58a6ff !important;
}

/* Waarschuwing */
.niet-gedig {
    background: #2d1f00;
    border: 1px solid #d29922;
    border-radius: 6px;
    padding: 0.8rem 1.2rem;
    color: #d29922;
    font-size: 0.85rem;
    margin: 0.5rem 0 1rem 0;
}

/* Sectie divider */
.sectie-header {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.7rem;
    color: #58a6ff;
    text-transform: uppercase;
    letter-spacing: 3px;
    border-bottom: 1px solid #21262d;
    padding-bottom: 0.5rem;
    margin: 1.5rem 0 1rem 0;
}

/* Verberg Streamlit branding */
#MainMenu, footer, header { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ── API configuratie ──────────────────────────────────────────────────────────
RP_API_KEY = "085ebb90bd31d7ce9a6c3ebfb40745e5"
RP_BASE    = "https://ruimte.omgevingswet.overheid.nl/ruimtelijke-plannen/api/opvragen/v4"
LS_BASE    = "https://api.pdok.nl/bzk/locatieserver/search/v3_1"

def rp_headers(met_body=False):
    h = {"X-Api-Key": RP_API_KEY, "Accept": "application/hal+json",
         "Content-Crs": "epsg:28992", "Accept-Crs": "epsg:28992"}
    if met_body: h["Content-Type"] = "application/json"
    return h

# ── DSO functies ──────────────────────────────────────────────────────────────
def adres_naar_rd(adres, log):
    params = {"q": adres, "fq": "type:adres", "rows": 100,
              "fl": "id,weergavenaam,centroide_rd"}
    resp = requests.get(f"{LS_BASE}/free", params=params, timeout=10)
    resp.raise_for_status()
    docs = resp.json().get("response", {}).get("docs", [])
    if not docs:
        raise ValueError(f"Adres niet gevonden: {adres}")

    adres_lower = adres.strip().lower()
    def is_match(w):
        wl = w.strip().lower()
        if wl == adres_lower: return True
        delen = wl.split(", ", 1)
        if len(delen) == 2:
            rest = delen[1].split(" ", 1)
            plaats = rest[1] if len(rest) == 2 else delen[1]
            if f"{delen[0]}, {plaats}" == adres_lower: return True
            if delen[0] == adres_lower: return True
        return False

    gezien, exacte = set(), []
    for d in docs:
        naam = d.get("weergavenaam", "")
        if is_match(naam) and naam not in gezien:
            gezien.add(naam); exacte.append(d)
    doc = exacte[0] if exacte else docs[0]
    log.append(f"  📍 Gevonden    : {doc.get('weergavenaam','—')}")

    rd = doc.get("centroide_rd","").replace("POINT(","").replace(")","").split()
    x, y = float(rd[0]), float(rd[1])
    log.append(f"  📐 Coördinaten : X={x:.0f}, Y={y:.0f}")

    kadastrale = "—"
    adres_id = doc.get("id","")
    if adres_id:
        r2 = requests.get(f"{LS_BASE}/lookup",
                          params={"id": adres_id, "fl": "gekoppeld_perceel"}, timeout=10)
        if r2.ok:
            ldocs = r2.json().get("response",{}).get("docs",[])
            if ldocs:
                percelen = ldocs[0].get("gekoppeld_perceel",[])
                if percelen: kadastrale = ", ".join(percelen)
    log.append(f"  🏠 Kadaster    : {kadastrale}")
    return {"x": x, "y": y, "weergavenaam": doc.get("weergavenaam", adres),
            "kadastrale_aanduiding": kadastrale}

def is_parapluplan(plan):
    if plan.get("isParapluplan") is True: return True
    naam = plan.get("naam","").lower()
    kws = ["paraplu","mantelzorg","parkeer","kruimelgeval","wooneenheid",
           "bed and breakfast","datacenter","darkstore","terrasregel",
           "terrassen","detailhandel","reclame","tam-omgevingsplan",
           "voorbereidingsbesluit","herziening"]
    return any(k in naam for k in kws)

def haal_bestemmingsplan(x, y, log):
    body = {"_geo": {"intersects": {"type": "Point", "coordinates": [x, y]}}}
    r = requests.post(f"{RP_BASE}/plannen/_zoek", headers=rp_headers(True),
                      params={"planType":"bestemmingsplan","planStatus":"vigerend",
                              "page":0,"pageSize":10},
                      json=body, timeout=15)
    r.raise_for_status()
    plannen = r.json().get("_embedded",{}).get("plannen",[])
    if not plannen: return None
    moeder = [p for p in plannen if not is_parapluplan(p)] or plannen
    plan = sorted(moeder, key=lambda p:(p.get("planstatusInfo") or {}).get("datum",""), reverse=True)[0]
    plan_id = plan.get("id","—")
    naam    = plan.get("naam","—")
    datum   = (plan.get("planstatusInfo") or {}).get("datum","—")
    log.append(f"  📋 Plan        : {naam}")
    log.append(f"  📅 Datum       : {datum}")
    hyperlink = f"https://www.ruimtelijkeplannen.nl/viewer/viewer?planidn={plan_id}"

    relevante_types = {t: [] for t in ["bestemmingsplan","omgevingsplan",
                                        "voorbereidingsbesluit","beheersverordening","inpassingsplan"]}
    for pt in ["omgevingsplan","voorbereidingsbesluit","beheersverordening","inpassingsplan"]:
        rx = requests.post(f"{RP_BASE}/plannen/_zoek", headers=rp_headers(True),
                           params={"planType":pt,"planStatus":"vigerend","page":0,"pageSize":5},
                           json=body, timeout=10)
        if rx.ok:
            for p in (rx.json().get("_embedded") or {}).get("plannen",[]):
                relevante_types[pt].append({"naam":p.get("naam","—"),
                    "datum":(p.get("planstatusInfo") or {}).get("datum","—"),
                    "id":p.get("id","—")})
    for p in plannen:
        relevante_types["bestemmingsplan"].append({"naam":p.get("naam","—"),
            "datum":(p.get("planstatusInfo") or {}).get("datum","—"),
            "id":p.get("id","—"),"paraplu":is_parapluplan(p)})

    alle_moeder = [{"id":p.get("id","—"),"naam":p.get("naam","—"),
                    "datum":(p.get("planstatusInfo") or {}).get("datum","—")}
                   for p in sorted(moeder, key=lambda p:(p.get("planstatusInfo") or {}).get("datum",""), reverse=True)]
    return {"id":plan_id,"naam":naam,"datum":datum,"hyperlink":hyperlink,
            "planenoverzicht":relevante_types,"alle_moederplannen":alle_moeder}

def haal_bestemmingsvlak(plan_id, x, y, log):
    body = {"_geo": {"intersects": {"type":"Point","coordinates":[x,y]}}}
    r = requests.post(f"{RP_BASE}/plannen/{plan_id}/bestemmingsvlakken/_zoek",
                      headers=rp_headers(True), params={"pageSize":10},
                      json=body, timeout=15)
    if not r.ok: return None
    vlakken = (r.json().get("_embedded") or {}).get("bestemmingsvlakken",[])
    if not vlakken: return None
    enkelvlak = [v for v in vlakken if v.get("type") == "enkelbestemming"]
    vlak = enkelvlak[0] if enkelvlak else vlakken[0]
    dubbel = [{"naam":v.get("naam","—"),"artikelnummer":v.get("artikelnummer","—"),"id":v.get("id","—")}
              for v in vlakken if v.get("type") == "dubbelbestemming"]
    log.append(f"  🗺️  Bestemming  : {vlak.get('naam','—')} ({vlak.get('type','—')})")
    return {"id":vlak.get("id","—"),"naam":vlak.get("naam","—"),
            "type":vlak.get("type","—"),"links":vlak.get("_links",{}),"dubbelbestemmingen":dubbel}

def haal_functieaanduidingen(vlak_links, log):
    fa_links = (vlak_links or {}).get("functieaanduidingen",[])
    if not fa_links: return []
    resultaat, seen = [], set()
    for link in (fa_links if isinstance(fa_links, list) else [fa_links]):
        href = link.get("href") if isinstance(link, dict) else None
        if not href: continue
        r = requests.get(href, headers=rp_headers(), timeout=15)
        if r.ok:
            naam = r.json().get("naam","—")
            if naam not in seen: seen.add(naam); resultaat.append(naam)
    if resultaat: log.append(f"  🏷️  Functieaand.: {', '.join(resultaat)}")
    else: log.append("  🏷️  Functieaand.: geen")
    return resultaat

def haal_maatvoeringen(plan_id, vlak_id, x, y, vlak_links, log):
    resultaat = []
    bv_links = (vlak_links or {}).get("bouwvlakken",[])
    if bv_links:
        for bv in bv_links:
            href = bv.get("href") if isinstance(bv, dict) else None
            if not href: continue
            bv_id = href.rstrip("/").split("/")[-1]
            r = requests.get(f"{RP_BASE}/plannen/{plan_id}/maatvoeringen",
                             headers=rp_headers(), params={"bouwvlak":bv_id,"pageSize":20}, timeout=15)
            if r.ok:
                for item in (r.json().get("_embedded") or {}).get("maatvoeringen",[]):
                    for o in item.get("omvang",[]):
                        resultaat.append({"naam":o.get("naam","—"),"waarde":o.get("waarde","—"),"eenheid":"m"})
                if resultaat: return resultaat
    body = {"_geo":{"intersects":{"type":"Point","coordinates":[x,y]}}}
    r = requests.post(f"{RP_BASE}/plannen/{plan_id}/maatvoeringen/_zoek",
                      headers=rp_headers(True), json=body, params={"pageSize":20}, timeout=15)
    if r.ok:
        for item in (r.json().get("_embedded") or {}).get("maatvoeringen",[]):
            for o in item.get("omvang",[item]):
                naam   = o.get("naam", item.get("naam","—"))
                waarde = o.get("waarde", item.get("waarde","—"))
                resultaat.append({"naam":naam,"waarde":waarde,"eenheid":"m"})
    for m in resultaat:
        log.append(f"  📏 {m['naam']:<28}: {m['waarde']} m")
    if not resultaat: log.append("  📏 Maatvoeringen: geen gevonden")
    return resultaat

def haal_data(adres):
    log = []
    resultaat = {
        "adres": adres, "kadastrale_aanduiding": "—",
        "bestemmingsplan_naam": "—", "bestemmingsplan_datum": "—",
        "hyperlink": "—", "bestemming_perceel": "—", "bestemmingstype": "—",
        "functieaanduidingen": [], "dubbelbestemmingen": [],
        "bouwaanduidingen": [], "maatvoeringen": [],
        "planenoverzicht": {t: [] for t in ["bestemmingsplan","omgevingsplan",
                            "voorbereidingsbesluit","beheersverordening","inpassingsplan"]},
        "niet_gedigitaliseerd": False,
    }

    log.append("Stap 1 — Adres omzetten naar coördinaten")
    loc = adres_naar_rd(adres, log)
    x, y = loc["x"], loc["y"]
    resultaat["adres_gevonden"]        = loc["weergavenaam"]
    resultaat["kadastrale_aanduiding"] = loc["kadastrale_aanduiding"]

    log.append("\nStap 2 — Vigerend bestemmingsplan ophalen")
    plan = haal_bestemmingsplan(x, y, log)
    if not plan:
        log.append("  ✗ Geen plan gevonden")
        return resultaat, log

    resultaat.update({"bestemmingsplan_naam": plan["naam"],
                      "bestemmingsplan_datum": plan["datum"],
                      "hyperlink": plan["hyperlink"],
                      "planenoverzicht": plan.get("planenoverzicht",{})})

    log.append("\nStap 3 — Bestemmingsvlak ophalen")
    vlak = haal_bestemmingsvlak(plan["id"], x, y, log)
    if not vlak and "alle_moederplannen" in plan:
        for ouder in plan["alle_moederplannen"][1:]:
            log.append(f"  ↩ Probeer ouder plan: {ouder['naam']}")
            vlak = haal_bestemmingsvlak(ouder["id"], x, y, log)
            if vlak:
                resultaat["bestemmingsplan_naam"] = ouder["naam"]
                resultaat["hyperlink"] = f"https://www.ruimtelijkeplannen.nl/viewer/viewer?planidn={ouder['id']}"
                break
    if not vlak:
        log.append("  ⚠ Plan niet gedigitaliseerd in API")
        resultaat["bestemming_perceel"]    = "↗ niet beschikbaar via API"
        resultaat["niet_gedigitaliseerd"]  = True
        return resultaat, log

    resultaat["bestemming_perceel"]  = vlak["naam"]
    resultaat["bestemmingstype"]     = vlak["type"]
    resultaat["dubbelbestemmingen"]  = vlak.get("dubbelbestemmingen",[])

    log.append("\nStap 4 — Functieaanduidingen ophalen")
    resultaat["functieaanduidingen"] = haal_functieaanduidingen(vlak.get("links",{}), log)

    log.append("\nStap 5 — Bouwaanduidingen ophalen")
    body = {"_geo":{"intersects":{"type":"Point","coordinates":[x,y]}}}
    br = requests.post(f"{RP_BASE}/plannen/{plan['id']}/bouwaanduidingen/_zoek",
                       headers=rp_headers(True), json=body, params={"pageSize":20}, timeout=15)
    bouw = []
    if br.ok:
        for ba in (br.json().get("_embedded") or {}).get("bouwaanduidingen",[]):
            bouw.append({"naam":ba.get("naam","—"),"artikelnummer":ba.get("artikelnummer","—")})
    resultaat["bouwaanduidingen"] = bouw
    log.append(f"  🔨 Bouwaand.   : {', '.join(b['naam'] for b in bouw) or 'geen'}")

    log.append("\nStap 6 — Maatvoeringen ophalen")
    resultaat["maatvoeringen"] = haal_maatvoeringen(plan["id"], vlak["id"], x, y, vlak.get("links",{}), log)

    log.append("\n✅ Alle data opgehaald!")
    return resultaat, log

# ── Word-document generator ───────────────────────────────────────────────────
def genereer_docx(data) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLAUW     = RGBColor(0x1F, 0x4E, 0x79)
    WIT       = RGBColor(0xFF, 0xFF, 0xFF)
    LB        = "D6E4F0"
    ROOD      = "FCE4D6"
    GEEL      = "FFF2CC"
    GRIJS     = "F2F2F2"
    BLAUW_HEX = "1F4E79"

    def set_bg(cell, hex):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex)
        tcPr.append(shd)

    def set_border(cell):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tblB=OxmlElement("w:tcBorders")
        for z in ("top","left","bottom","right"):
            b=OxmlElement(f"w:{z}"); b.set(qn("w:val"),"single")
            b.set(qn("w:sz"),"4"); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"AAAAAA")
            tblB.append(b)
        tcPr.append(tblB)

    def cel(cell, tekst, vet=False, cur=False, kleur=None, pt=9):
        cell.text=""; p=cell.paragraphs[0]; r=p.add_run(str(tekst))
        r.bold=vet; r.italic=cur; r.font.size=Pt(pt); r.font.name="Arial"
        if kleur: r.font.color.rgb=kleur
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)

    def marge(cell):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar")
        for z,v in [("top","60"),("bottom","60"),("left","120"),("right","120")]:
            m=OxmlElement(f"w:{z}"); m.set(qn("w:w"),v); m.set(qn("w:type"),"dxa"); mar.append(m)
        tcPr.append(mar)

    def header(doc, tekst):
        p=doc.add_paragraph(); r=p.add_run(tekst)
        r.bold=True; r.font.size=Pt(11); r.font.name="Arial"; r.font.color.rgb=WIT
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),BLAUW_HEX)
        pPr.append(shd); ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"120"); pPr.append(ind)

    def tabel(doc):
        t=doc.add_table(rows=0,cols=2); t.style="Table Grid"; t.autofit=False
        t.columns[0].width=Cm(5.5); t.columns[1].width=Cm(10.5); return t

    def rij(t, label, waarde, kleur="wit"):
        r=t.add_row(); cl,cw=r.cells[0],r.cells[1]
        set_border(cl); set_border(cw); set_bg(cl,GRIJS)
        bg={"blauw":LB,"rood":ROOD,"geel":GEEL}.get(kleur,"FFFFFF")
        set_bg(cw,bg); cur=kleur in ("blauw","rood")
        cel(cl,label,vet=True,pt=9); marge(cl)
        cel(cw,str(waarde) if waarde else "",cur=cur,pt=9); marge(cw)

    def auto(waarde, leeg="geen"):
        if waarde and "niet beschikbaar" in str(waarde).lower(): return waarde,"geel"
        v=waarde if waarde and waarde not in ("—","geen","") else leeg
        return v,"blauw"

    # Data uitlezen
    adres      = data.get("adres_gevonden") or data.get("adres","—")
    kadaster   = data.get("kadastrale_aanduiding","—")
    bp_naam    = data.get("bestemmingsplan_naam","—")
    bp_datum   = data.get("bestemmingsplan_datum","—")
    hyperlink  = data.get("hyperlink","—")
    bestemming = data.get("bestemming_perceel","—")
    maatvoeringen = data.get("maatvoeringen",[])
    niet_gedig = data.get("niet_gedigitaliseerd",False)
    NA         = "↗ niet beschikbaar via API — zie hyperlink plan"

    functie_str = ", ".join(data.get("functieaanduidingen",[])) or "geen"
    dubbel_str  = ", ".join(d["naam"] for d in data.get("dubbelbestemmingen",[])) or "geen"
    bouw_str    = ", ".join(b["naam"] for b in data.get("bouwaanduidingen",[])) or "geen"

    if niet_gedig:
        bestemming=functie_str=dubbel_str=bouw_str=NA

    def maatv(zoektermen):
        for z in zoektermen:
            for m in maatvoeringen:
                if z in m.get("naam","").lower():
                    return f"{m['waarde']} {m.get('eenheid','')}".strip()
        return "—"

    bouwhoogte = NA if niet_gedig else maatv(["bouwhoogte"])
    goothoogte = NA if niet_gedig else maatv(["goothoogte"])
    opp        = NA if niet_gedig else maatv(["oppervlakte","bebouwd","bouwperceel"])

    adres_delen = adres.split(", ") if ", " in adres else [adres]
    straat_hnr  = adres_delen[0]
    woonplaats  = adres_delen[-1].split(" ",1)[-1] if len(adres_delen)>=2 else "—"

    planenoverzicht = data.get("planenoverzicht",{})
    op_naam = "—"
    for p in planenoverzicht.get("omgevingsplan",[]): op_naam=p.get("naam","—"); break
    vbb_str = ", ".join(p.get("naam","—") for p in planenoverzicht.get("voorbereidingsbesluit",[])) or "—"
    parkeer_str = "—"
    for p in planenoverzicht.get("bestemmingsplan",[]):
        if "parkeer" in p.get("naam","").lower() and p.get("paraplu"):
            parkeer_str=p.get("naam","—"); break

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Cm(2)

    p=doc.add_paragraph()
    r=p.add_run("INTAKE TOETS")
    r.bold=True; r.font.size=Pt(18); r.font.color.rgb=BLAUW; r.font.name="Arial"
    p.paragraph_format.space_after=Pt(2)

    p2=doc.add_paragraph()
    r2=p2.add_run(
        f"Gegenereerd op {date.today().strftime('%d-%m-%Y')} via DSO Intake Toets Generator  |  "
        "Lichtblauw = automatisch  |  Rood = via PowerBrowser  |  Geel = handmatig invullen"
    )
    r2.font.size=Pt(7.5); r2.font.name="Arial"; r2.font.color.rgb=RGBColor(0x99,0x99,0x99)
    p2.paragraph_format.space_after=Pt(6)

    header(doc,"1. ALGEMENE ZAAKGEGEVENS")
    t=tabel(doc)
    for lbl,klr in [("Zaak-ID","geel"),("Kenmerk","geel"),("Ontvangstdatum","rood"),
                    ("DSO Verzoeknummer","rood"),("Casemanager","rood"),("Afdeling","rood"),
                    ("Datum adviesaanvraag RO","geel"),("Aanvraagtype","rood"),
                    ("Gerelateerde zaken","geel"),("Omschrijving bouwplan","rood"),("Bijzonderheden","rood")]:
        rij(t,lbl,"",klr)

    header(doc,"2. LOCATIE")
    t=tabel(doc)
    rij(t,"Straatnaam + huisnummer",straat_hnr,"blauw")
    rij(t,"Woonplaats",woonplaats,"blauw")
    rij(t,"Kadastrale aanduiding",kadaster,"blauw")

    header(doc,"3. ACTIVITEITEN")
    for act in ["Bouwactiviteit (technisch)","Bouwactiviteit (omgevingsplan)",
                "Omgevingsplanactiviteit (uitvoeren van een werk)",
                "Omgevingsplanactiviteit (slopen)","Omgevingsplanactiviteit (afwijken van de regels)",
                "Rijksmonumentenactiviteit (Bouwwerk)","Milieubelastende activiteit"]:
        p=doc.add_paragraph()
        r=p.add_run("☐ "+act); r.font.size=Pt(9); r.font.name="Arial"
        p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.space_before=p.paragraph_format.space_after=Pt(1)

    header(doc,"6. OMGEVINGSPLANTOETS")
    t=tabel(doc)
    rij(t,"Hyperlink Regels op de kaart",
        "https://omgevingswet.overheid.nl/regels-op-de-kaart/zoeken/locatie","blauw")

    if niet_gedig:
        p_w=doc.add_paragraph()
        pPr=p_w._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"FFD966")
        pPr.append(shd)
        rw=p_w.add_run(f"⚠  Dit plan ({bp_naam}) is niet volledig gedigitaliseerd. "
                        f"Gele velden handmatig opzoeken via: {hyperlink}")
        rw.bold=True; rw.font.size=Pt(8); rw.font.name="Arial"
        rw.font.color.rgb=RGBColor(0x7F,0x3F,0x00)

    omg_w = op_naam if op_naam!="—" else "↗ zie Regels op de kaart"
    rij(t,"Omgevingsplan",omg_w,"blauw")
    rij(t,"Bestemmingsplan",*auto(bp_naam))
    rij(t,"Datum vaststelling",*auto(bp_datum))
    rij(t,"Hyperlink plan",*auto(hyperlink))
    rij(t,"Bestemming perceel",*auto(bestemming))
    rij(t,"Dubbelbestemming",*auto(dubbel_str))
    rij(t,"(Functie)aanduiding",*auto(functie_str))
    rij(t,"Bouwaanduiding",*auto(bouw_str))
    rij(t,"Voorbereidingsbesluit",*auto(vbb_str,"geen"))
    rij(t,"Bestemmingsplan parkeren",*auto(parkeer_str,"geen"))
    rij(t,"Bebouwde oppervlakte",*auto(opp,"niet opgenomen in plan"))
    rij(t,"Maximale bouwhoogte",*auto(bouwhoogte,"niet opgenomen in plan"))
    rij(t,"Maximale goothoogte",*auto(goothoogte,"niet opgenomen in plan"))
    for m in maatvoeringen:
        if not any(k in m.get("naam","").lower() for k in ["bouwhoogte","goothoogte","oppervlakte"]):
            rij(t,m["naam"],f"{m['waarde']} {m.get('eenheid','')}".strip(),"blauw")

    header(doc,"7. CONCLUSIE OMGEVINGSPLANTOETS")
    t=tabel(doc)
    for lbl in ["Voldoet aan bestemmingsplan?","Voldoet aan afwijkingsregels?",
                "BOPA ja of nee?","Procedure","Anterieure overeenkomst?"]:
        rij(t,lbl,"","geel")

    header(doc,"8. TOETS OMGEVINGSPLAN")
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),GEEL)
    pPr.append(shd)
    plan_w=op_naam if op_naam!="—" else bp_naam
    for tekst,vet,cur,kleur in [
        ("Het perceel ",False,False,None),
        (f"{straat_hnr}, {woonplaats}",True,False,BLAUW),
        (" ligt binnen het plangebied van ",False,False,None),
        (f'"{plan_w}"',False,True,None),
        (". Het perceel heeft de bestemming ",False,False,None),
        (f"'{bestemming}'",True,False,None),
        (". De aanvraag is gesitueerd op deze bestemming.",False,False,None),
    ]:
        r=p.add_run(tekst); r.bold=vet; r.italic=cur
        r.font.size=Pt(9); r.font.name="Arial"
        if kleur: r.font.color.rgb=kleur

    p_voet=doc.add_paragraph()
    rv=p_voet.add_run(f"Gegenereerd op {date.today().strftime('%d-%m-%Y')} via DSO Intake Toets Generator v2.0")
    rv.font.size=Pt(7); rv.font.name="Arial"; rv.font.color.rgb=RGBColor(0x99,0x99,0x99)
    p_voet.paragraph_format.space_before=Pt(8)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── UI ────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="dso-header">
    <h1>🏛️ DSO Intake Toets Generator</h1>
    <p>Automatisch bestemmingsplandata ophalen en invullen in een Word-document</p>
</div>
""", unsafe_allow_html=True)

st.markdown('<span class="badge-auto">Lichtblauw = automatisch (DSO)</span>'
            '<span class="badge-pb">Rood = via PowerBrowser</span>'
            '<span class="badge-hand">Geel = handmatig invullen</span>',
            unsafe_allow_html=True)

st.markdown("---")

# Invoer
st.markdown('<div class="stap-label">Stap 1 — Voer het adres in</div>', unsafe_allow_html=True)
adres_input = st.text_input(
    label="Adres",
    placeholder="bijv. Kerkstraat 1, IJsselstein",
    label_visibility="collapsed"
)

col1, col2 = st.columns([2, 1])
with col1:
    zoek_knop = st.button("🔍  Data ophalen + Word genereren", type="primary")
with col2:
    if st.button("Testadres gebruiken"):
        st.session_state["testadres"] = "Prinsengracht 40A, Amsterdam"
        st.rerun()

if "testadres" in st.session_state:
    adres_input = st.session_state.pop("testadres")

# Verwerking
if zoek_knop and adres_input:
    st.markdown("---")
    st.markdown('<div class="stap-label">Stap 2 — Data ophalen via DSO API</div>', unsafe_allow_html=True)

    log_placeholder = st.empty()
    log_tekst = ""

    with st.spinner("Bezig met ophalen..."):
        try:
            data, log = haal_data(adres_input.strip())
            log_tekst = "\n".join(log)
        except Exception as e:
            st.error(f"❌ Fout: {e}")
            st.stop()

    log_placeholder.markdown(f'<div class="terminal">{log_tekst}</div>', unsafe_allow_html=True)

    # Resultaten tonen
    st.markdown("---")
    st.markdown('<div class="stap-label">Stap 3 — Gevonden gegevens</div>', unsafe_allow_html=True)

    if data.get("niet_gedigitaliseerd"):
        st.markdown(f"""<div class="niet-gedig">
            ⚠️ <strong>Dit plan is niet volledig gedigitaliseerd</strong> in de Ruimtelijke Plannen API.<br>
            Gele velden in het Word-document moeten handmatig worden ingevuld via:<br>
            <a href="{data.get('hyperlink','#')}" target="_blank">{data.get('hyperlink','—')}</a>
        </div>""", unsafe_allow_html=True)

    def kaart(label, waarde, auto=True):
        klasse = "auto" if auto and waarde and waarde not in ("—","geen") else ("leeg" if not waarde or waarde=="—" else "")
        w = waarde if waarde and waarde != "—" else "—"
        return f"""<div class="result-card">
            <div class="label">{label}</div>
            <div class="waarde {klasse}">{w}</div>
        </div>"""

    col1, col2 = st.columns(2)
    with col1:
        st.markdown(kaart("Gevonden adres", data.get("adres_gevonden","—")), unsafe_allow_html=True)
        st.markdown(kaart("Kadastrale aanduiding", data.get("kadastrale_aanduiding","—")), unsafe_allow_html=True)
        st.markdown(kaart("Bestemmingsplan", data.get("bestemmingsplan_naam","—")), unsafe_allow_html=True)
        st.markdown(kaart("Datum vaststelling", data.get("bestemmingsplan_datum","—")), unsafe_allow_html=True)
    with col2:
        st.markdown(kaart("Bestemming perceel", data.get("bestemming_perceel","—")), unsafe_allow_html=True)
        st.markdown(kaart("Bestemmingstype", data.get("bestemmingstype","—")), unsafe_allow_html=True)
        functie = ", ".join(data.get("functieaanduidingen",[])) or "geen"
        dubbel  = ", ".join(d["naam"] for d in data.get("dubbelbestemmingen",[])) or "geen"
        st.markdown(kaart("Functieaanduiding", functie), unsafe_allow_html=True)
        st.markdown(kaart("Dubbelbestemming", dubbel), unsafe_allow_html=True)

    if data.get("maatvoeringen"):
        st.markdown('<div class="sectie-header">Maatvoeringen</div>', unsafe_allow_html=True)
        cols = st.columns(3)
        for i, m in enumerate(data["maatvoeringen"]):
            with cols[i % 3]:
                st.markdown(kaart(m["naam"], f"{m['waarde']} {m.get('eenheid','')}".strip()), unsafe_allow_html=True)

    # Word genereren
    st.markdown("---")
    st.markdown('<div class="stap-label">Stap 4 — Download Word-document</div>', unsafe_allow_html=True)

    with st.spinner("Word-document genereren..."):
        buf = genereer_docx(data)

    adres_kort = adres_input.split(",")[0].replace(" ","_").replace("/","-")[:25]
    bestandsnaam = f"Intake_toets_{adres_kort}_{date.today().strftime('%Y%m%d')}.docx"

    st.download_button(
        label="📄  Download Intake Toets (.docx)",
        data=buf,
        file_name=bestandsnaam,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.success(f"✅ Klaar! Klik op de knop om **{bestandsnaam}** te downloaden.")

elif zoek_knop and not adres_input:
    st.warning("⚠️ Vul eerst een adres in.")
